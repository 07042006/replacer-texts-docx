import re
from docx import Document

def docx_replace_regex(doc_obj, regex, replace, regexintext):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

oldwords = ["SC", "SE", "PBROOL", "SCCOND", "PBCOND"] # palavaras antigas
newwords = ["PB", "PB", "SCROOL", "SECOND", "SECOND"] # palavras novas (substituirão)

filename = "filetest.docx"
doc = Document(filename)
counter = 0

for replecingwords in oldwords:
    regextext = str(oldwords[counter])
    regex1 = re.compile(str(oldwords[counter]))
    replace1 = str(newwords[counter])

    docx_replace_regex(doc, regex1 , replace1, regextext)
    counter = counter+1

doc.save("updated_fileteste.docx") # save filedocx
